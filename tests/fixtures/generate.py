"""Regenerate the binary test fixtures in this directory.

Run with: uv run python tests/fixtures/generate.py

The fixtures are committed so the test suite needs no generation step; this
script exists so they stay reproducible and auditable. Everything is synthetic
— no real documents, no sensitive data. PDFs are written as uncompressed
PDF 1.4 by hand to avoid pulling in a PDF-authoring dependency.
"""

from pathlib import Path

FIXTURES = Path(__file__).parent

_PAGE_W, _PAGE_H = 612, 792


def _escape(text: str) -> str:
    return text.replace("\\", r"\\").replace("(", r"\(").replace(")", r"\)")


def _text_ops(
    lines: list[str], x: int = 72, y: int = 720, size: int = 12, bold: bool = False
) -> str:
    """Build a PDF text object drawing one line per entry, top-down."""
    font = "/F2" if bold else "/F1"
    ops = [f"BT {font} {size} Tf {x} {y} Td {size + 6} TL"]
    for i, line in enumerate(lines):
        if i:
            ops.append("T*")
        ops.append(f"({_escape(line)}) Tj")
    ops.append("ET")
    return "\n".join(ops)


def _rule_ops(rules: list[tuple[int, int, int, int]]) -> str:
    """Build stroked line segments (x1 y1 x2 y2), used to rule table borders."""
    if not rules:
        return ""
    ops = ["0.5 w"]
    for x1, y1, x2, y2 in rules:
        ops.append(f"{x1} {y1} m {x2} {y2} l S")
    return "\n".join(ops)


_YELLOW = "1 1 0"
_GREEN = "0.3 0.9 0.3"


def _fill_rect_ops(rects: list[tuple[float, float, float, float, str]]) -> str:
    """Filled rectangles (x, y, w, h, color), emitted BEFORE the text ops so
    the fill sits behind the glyphs — exactly how Word exports highlighting."""
    ops = []
    for x, y, w, h, color in rects:
        ops.append(f"{color} rg\n{x} {y} {w} {h} re\nf")
    return "\n".join(ops)


def _build_pdf(pages: list[str]) -> bytes:
    """Assemble page content streams into a minimal two-font PDF."""
    n = len(pages)
    font_id = 3 + 2 * n
    bold_id = font_id + 1
    objects: dict[int, str] = {
        1: "<< /Type /Catalog /Pages 2 0 R >>",
        2: (
            "<< /Type /Pages /Count {n} /Kids [{kids}] >>".format(
                n=n, kids=" ".join(f"{3 + 2 * i} 0 R" for i in range(n))
            )
        ),
        font_id: "<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        bold_id: "<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica-Bold >>",
    }
    for i, content in enumerate(pages):
        page_id, stream_id = 3 + 2 * i, 4 + 2 * i
        objects[page_id] = (
            f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 {_PAGE_W} {_PAGE_H}] "
            f"/Resources << /Font << /F1 {font_id} 0 R /F2 {bold_id} 0 R >> >> "
            f"/Contents {stream_id} 0 R >>"
        )
        objects[stream_id] = (
            f"<< /Length {len(content.encode())} >>\nstream\n{content}\nendstream"
        )

    out = bytearray(b"%PDF-1.4\n")
    offsets: dict[int, int] = {}
    for obj_id in sorted(objects):
        offsets[obj_id] = len(out)
        out += f"{obj_id} 0 obj\n{objects[obj_id]}\nendobj\n".encode()

    xref_pos = len(out)
    max_id = max(objects)
    out += f"xref\n0 {max_id + 1}\n".encode()
    out += b"0000000000 65535 f \n"
    for obj_id in range(1, max_id + 1):
        out += f"{offsets[obj_id]:010d} 00000 n \n".encode()
    out += (
        f"trailer\n<< /Size {max_id + 1} /Root 1 0 R >>\nstartxref\n{xref_pos}\n%%EOF\n"
    ).encode()
    return bytes(out)


def make_native_text_pdf() -> None:
    """Case A — a PDF whose text is embedded, no OCR needed."""
    content = "\n".join(
        [
            _text_ops(["Deployment Runbook"], y=724, size=17, bold=True),
            _text_ops(
                [
                    "The ingestion service reads documents from the corpus root and",
                    "writes chunks to the vector store. Restart it with systemctl and",
                    "wait for the readiness probe to report a healthy status before",
                    "sending production traffic back to the instance.",
                    "",
                    "The rollback token for the staging cluster is RBK-7781-ZULU.",
                ],
                y=688,
                size=11,
            ),
        ]
    )
    (FIXTURES / "native_text.pdf").write_bytes(_build_pdf([content]))


def make_multipage_pdf() -> None:
    """Case D — three pages with page-unique sentinels for page_number checks.

    Each page gets its own bold heading and a full body paragraph. The heading
    gives every page distinct chunk metadata and the body makes each page large
    enough that the chunker does not merge pages into a single chunk — without
    that, page provenance would be untestable.
    """
    pages = []
    chapters = [
        ("Ingestion Subsystem", "ALPHA-PAGE-ONE"),
        ("Retrieval Subsystem", "BRAVO-PAGE-TWO"),
        ("Evaluation Subsystem", "CHARLIE-PAGE-THREE"),
    ]
    for index, (title, sentinel) in enumerate(chapters, start=1):
        body = [
            f"This chapter documents the {title.lower()} of the platform and the",
            "operational duties that belong to the team that owns it. The service",
            "reads its configuration at startup and refuses to boot when a required",
            "value is missing, which keeps a partial deployment from serving traffic.",
            "Operators are expected to review the dashboards after every release and",
            "record any anomaly in the shared incident log before closing the change.",
            "Routine maintenance happens during the weekly window agreed with the",
            "platform group, and emergency work follows the on-call escalation path.",
            "",
            f"The verification marker for this chapter is {sentinel}, and that marker",
            "appears on no other page of this document.",
        ]
        pages.append(
            "\n".join(
                [
                    _text_ops([f"Chapter {index}: {title}"], y=724, size=17, bold=True),
                    _text_ops(body, y=688, size=11),
                ]
            )
        )
    (FIXTURES / "multipage.pdf").write_bytes(_build_pdf(pages))


def make_table_pdf() -> None:
    """Case C — a ruled table so TableFormer recovers row/column structure."""
    rows = [
        ("Region", "Quarter", "Revenue"),
        ("North", "Q1", "118400"),
        ("North", "Q2", "126950"),
        ("South", "Q1", "94300"),
        ("South", "Q2", "101275"),
    ]
    cols_x = [90, 240, 390]
    top_y = 660
    row_h = 26

    ops = [_text_ops(["Quarterly Revenue Report"], y=724, size=17, bold=True)]
    for r, row in enumerate(rows):
        y = top_y - r * row_h
        for c, cell in enumerate(row):
            ops.append(_text_ops([cell], x=cols_x[c] + 6, y=y + 8, size=11))

    rules: list[tuple[int, int, int, int]] = []
    bottom_y = top_y - len(rows) * row_h + row_h - 8
    for r in range(len(rows) + 1):
        y = top_y + 22 - r * row_h
        rules.append((90, y, 520, y))
    for x in [90, 240, 390, 520]:
        rules.append((x, top_y + 22, x, bottom_y - 4))
    ops.append(_rule_ops(rules))

    (FIXTURES / "table.pdf").write_bytes(_build_pdf(["\n".join(ops)]))


def _render_text_image(lines: list[str], width: int = 1400, height: int = 500):
    from PIL import Image, ImageDraw, ImageFont

    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    font = ImageFont.load_default(size=54)
    y = 60
    for line in lines:
        draw.text((60, y), line, fill="black", font=font)
        y += 84
    return img


def make_scanned_pdf() -> None:
    """Case B — an image-only PDF, the OCR trigger case."""
    img = _render_text_image(
        [
            "SCANNED MAINTENANCE NOTICE",
            "Generator unit GX-4400 is offline.",
            "Service ticket number 88213 is open.",
        ]
    )
    img.save(FIXTURES / "scanned.pdf", "PDF", resolution=150.0)


def make_image_png() -> None:
    """Case E — a standalone raster image routed through the IMAGE backend."""
    img = _render_text_image(
        [
            "SAFETY PLACARD",
            "Wear eye protection in bay 7.",
        ],
        height=300,
    )
    img.save(FIXTURES / "image.png")


def make_docx() -> None:
    """Case F — an Office document with headings, a paragraph and a table."""
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_heading("Vendor Agreement", level=1)
    doc.add_heading("Payment Terms", level=2)
    doc.add_paragraph(
        "Invoices are settled within thirty days of receipt. "
        "The escalation contact is the account manager on record."
    )
    table = doc.add_table(rows=3, cols=2)
    data = [("Milestone", "Amount"), ("Kickoff", "15000"), ("Delivery", "42500")]
    for row, values in zip(table.rows, data):
        for cell, value in zip(row.cells, values):
            cell.text = value
    doc.save(FIXTURES / "sample.docx")


def make_highlighted_pdf() -> None:
    """Case G — Word-style text highlights: filled rects behind native text.

    Not annotations: the reference corpus (.docx -> PDF exports) encodes its
    yellow/green marks as content-stream fills, which is exactly what the
    emphasis extractor's rect path reads. Page 1 carries a CAPS yellow heading,
    a green state mark, a yellow-marked CPF (redaction interplay) and a 6x4pt
    decorative sliver the size filter must discard; page 2 exists so the
    page-mapping of spans is testable.
    """
    size = 12

    def line_rect(
        x: float, y: float, text: str, color: str, size: int = 12
    ) -> tuple[float, float, float, float, str]:
        # Generous width: the crop only ever sees this line, so covering a bit
        # of trailing whitespace changes nothing.
        return (x - 3, y - 3, len(text) * 0.75 * size, size + 5, color)

    heading = "PRAZO DE QUITACAO DO BOLETO"
    green_line = "enviar o boleto ao cliente"
    cpf_line = "Cpf do titular: 529.982.247-25"
    page1 = "\n".join(
        [
            _fill_rect_ops(
                [
                    line_rect(72, 680, heading, _YELLOW, size=14),
                    line_rect(72, 620, green_line, _GREEN),
                    line_rect(72, 590, cpf_line, _YELLOW),
                    (500, 400, 6, 4, _YELLOW),  # decorative sliver: filtered
                ]
            ),
            _text_ops(["Manual de atendimento ao cliente."], y=720, size=size),
            _text_ops([heading], y=680, size=14, bold=True),
            _text_ops(
                ["O procedimento abaixo descreve a rotina de cobranca."],
                y=650,
                size=size,
            ),
            _text_ops([green_line], y=620, size=size),
            _text_ops([cpf_line], y=590, size=size),
            _text_ops(["Encerrar o atendimento apos a confirmacao."], y=560, size=size),
        ]
    )
    reemissao = "REEMISSAO EM ATE 2 DIAS UTEIS"
    page2 = "\n".join(
        [
            _fill_rect_ops([line_rect(72, 700, reemissao, _YELLOW, size=14)]),
            _text_ops([reemissao], y=700, size=14, bold=True),
            _text_ops(
                ["A reemissao segue a fila padrao de atendimento."], y=660, size=size
            ),
        ]
    )
    (FIXTURES / "highlighted.pdf").write_bytes(_build_pdf([page1, page2]))


def main() -> None:
    make_native_text_pdf()
    make_multipage_pdf()
    make_table_pdf()
    make_scanned_pdf()
    make_image_png()
    make_docx()
    make_highlighted_pdf()
    for path in sorted(FIXTURES.iterdir()):
        if path.name != "generate.py":
            print(f"{path.name:24} {path.stat().st_size:>8} bytes")


if __name__ == "__main__":
    main()
